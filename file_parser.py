"""
file_parser.py
--------------
Extracts plain text from uploaded resume/document files.

Supported formats
-----------------
  .txt / .text   – read as UTF-8 text
  .pdf           – extract with pdfplumber (digital PDFs)
  .docx          – extract with python-docx
  .doc           – best-effort: try python-docx; return error hint on failure
  .png / .jpg / .jpeg / .webp / .bmp / .tiff / .gif
                 – OCR via Tesseract (pytesseract + Pillow)

All public functions return a plain str.
On any extraction failure the functions return an error-hint string rather
than raising so callers can surface a helpful message to the user.
"""

from __future__ import annotations

import io
import logging
import os

logger = logging.getLogger(__name__)

# ── Allowed extensions ────────────────────────────────────────────────────────

ALLOWED_EXTENSIONS = {
    ".txt", ".text",
    ".pdf",
    ".docx", ".doc",
    ".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".gif",
}

MAX_BYTES = 10 * 1024 * 1024  # 10 MB hard limit


def allowed_file(filename: str) -> bool:
    """Return True when *filename* has a supported extension."""
    if not filename:
        return False
    ext = os.path.splitext(filename.lower())[1]
    return ext in ALLOWED_EXTENSIONS


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Dispatch to the appropriate extractor based on *filename* extension.

    Parameters
    ----------
    file_bytes : raw bytes of the uploaded file
    filename   : original filename (used to determine format)

    Returns
    -------
    str – extracted plain text, possibly empty or an error message starting
          with '[Error]' so callers can distinguish extraction failures.
    """
    if len(file_bytes) > MAX_BYTES:
        return "[Error] File is too large. Please upload a file smaller than 10 MB."

    ext = os.path.splitext(filename.lower())[1]

    if ext in (".txt", ".text"):
        return _extract_txt(file_bytes)
    elif ext == ".pdf":
        return _extract_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_bytes, ext)
    elif ext in (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".gif"):
        return _extract_image(file_bytes)
    else:
        return f"[Error] Unsupported file type '{ext}'. Please upload a .txt, .pdf, .docx, or image file."


# ── Format-specific extractors ────────────────────────────────────────────────

def _extract_txt(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return data.decode(enc).strip()
        except (UnicodeDecodeError, ValueError):
            continue
    return "[Error] Could not decode text file. Please save it as UTF-8 and try again."


def _extract_pdf(data: bytes) -> str:
    try:
        import pdfplumber
        text_parts: list[str] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        result = "\n".join(text_parts).strip()
        if result:
            return result
        # PDF has no selectable text — fall through to image OCR
        logger.info("PDF has no embedded text; attempting page-image OCR.")
        return _extract_pdf_via_ocr(data)
    except ImportError:
        return "[Error] PDF parsing library not installed. Run: pip install pdfplumber"
    except Exception as exc:
        logger.error("PDF extraction failed: %s", exc)
        return f"[Error] Could not read PDF: {exc}"


def _extract_pdf_via_ocr(data: bytes) -> str:
    """OCR fallback for scanned/image-only PDFs using pypdfium2 + Tesseract."""
    try:
        import pypdfium2 as pdfium
        from PIL import Image
        import pytesseract

        pdf = pdfium.PdfDocument(data)
        texts: list[str] = []
        for page_index in range(len(pdf)):
            page   = pdf[page_index]
            bitmap = page.render(scale=2)        # 2× scale for better OCR
            pil_img = bitmap.to_pil()
            page_text = pytesseract.image_to_string(pil_img)
            if page_text.strip():
                texts.append(page_text)
        pdf.close()
        result = "\n".join(texts).strip()
        return result if result else "[Error] No text could be extracted from this PDF (scanned image with no recognisable text)."
    except Exception as exc:
        logger.error("PDF OCR failed: %s", exc)
        return "[Error] This PDF appears to be a scanned image and OCR failed. Please paste your resume text manually."


def _extract_docx(data: bytes, ext: str) -> str:
    try:
        from docx import Document
        doc   = Document(io.BytesIO(data))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
        result = "\n".join(lines).strip()
        return result if result else "[Error] The Word document appears to be empty."
    except ImportError:
        return "[Error] Word document library not installed. Run: pip install python-docx"
    except Exception as exc:
        logger.error("DOCX extraction failed: %s", exc)
        if ext == ".doc":
            return (
                "[Error] Legacy .doc format is not fully supported. "
                "Please save your document as .docx (Word 2007+) and re-upload, "
                "or paste the text directly."
            )
        return f"[Error] Could not read Word document: {exc}"


def _extract_image(data: bytes) -> str:
    """OCR an image file (PNG, JPG, etc.) using Tesseract."""
    try:
        from PIL import Image, ImageEnhance, ImageFilter
        import pytesseract

        img = Image.open(io.BytesIO(data)).convert("RGB")

        # Pre-processing: increase contrast & sharpen for better OCR accuracy
        img = ImageEnhance.Contrast(img).enhance(1.5)
        img = img.filter(ImageFilter.SHARPEN)

        # Scale up small images
        w, h = img.size
        if w < 1000:
            scale = 1000 / w
            img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

        text = pytesseract.image_to_string(img, config="--oem 3 --psm 6")
        result = text.strip()
        if result:
            return result
        return "[Error] No text could be extracted from this image. Make sure the image is clear and text is legible."
    except ImportError:
        return (
            "[Error] OCR libraries not installed. Run: pip install Pillow pytesseract\n"
            "Also install Tesseract-OCR from https://github.com/tesseract-ocr/tesseract"
        )
    except Exception as exc:
        logger.error("Image OCR failed: %s", exc)
        return f"[Error] Could not extract text from image: {exc}"


# ── Human-readable format list (for UI hints) ────────────────────────────────

ACCEPTED_FORMATS_LABEL = ".pdf, .docx, .doc, .txt, .png, .jpg, .jpeg, .webp, .bmp, .tiff"
ACCEPTED_MIME_TYPES    = (
    "application/pdf,"
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document,"
    "application/msword,"
    "text/plain,"
    "image/png,image/jpeg,image/webp,image/bmp,image/tiff,image/gif"
)
